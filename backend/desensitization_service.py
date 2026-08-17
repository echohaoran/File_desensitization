"""
脱敏服务模块
集成 Microsoft Presidio 和中国百家姓库进行 PII 检测和脱敏
"""

import re
import json
from typing import List, Dict, Any, Tuple, Optional
from datetime import datetime

# Microsoft Presidio 为可选增强能力；缺失时保留规则与中文姓名识别。
try:
    from presidio_analyzer import AnalyzerEngine
    from presidio_anonymizer import AnonymizerEngine
    PRESIDIO_AVAILABLE = True
except ImportError:
    AnalyzerEngine = AnonymizerEngine = None
    PRESIDIO_AVAILABLE = False

# 中文分词
import jieba

# PDF 处理
import PyPDF2
import io
import ipaddress

# Word 文档处理
from docx import Document

# Excel 处理
from openpyxl import load_workbook

# 尝试导入 spacy，如果失败则使用备用方案
try:
    import spacy
    SPACY_AVAILABLE = True
except ImportError:
    SPACY_AVAILABLE = False
    print("警告：spacy 未安装，将使用备用的 Presidio 分析器")


class DesensitizationService:
    """脱敏服务类"""
    
    def __init__(self):
        """初始化脱敏服务"""
        # 初始化 Presidio 分析器
        self.analyzer = None
        self.anonymizer = None
        
        try:
            # 检查 spacy 模型是否可用
            import spacy
            try:
                spacy.load("en_core_web_sm")
                self.analyzer = AnalyzerEngine()
                self.anonymizer = AnonymizerEngine()
                print("✓ Presidio 分析器初始化成功")
            except OSError:
                print("警告：spacy 模型未安装，将仅使用正则表达式和百家姓库进行检测")
                print("提示：如需 Presidio 功能，请运行 'python -m spacy download en_core_web_sm'")
        except ImportError:
            print("警告：spacy 未安装，将仅使用正则表达式和百家姓库进行检测")
        
        # 中国百家姓（常见姓氏）
        self.chinese_surnames = self._load_chinese_surnames()
        
        # 敏感信息模式
        self.sensitive_patterns = self._load_sensitive_patterns()
        
        # 统计信息
        self.stats = {
            "total_detections": 0,
            "by_type": {}
        }
    
    def _load_chinese_surnames(self) -> set:
        """加载中国百家姓"""
        # 常见姓氏列表（前100个）
        common_surnames = {
            "赵", "钱", "孙", "李", "周", "吴", "郑", "王", "冯", "陈",
            "褚", "卫", "蒋", "沈", "韩", "杨", "朱", "秦", "尤", "许",
            "何", "吕", "施", "张", "孔", "曹", "严", "华", "金", "魏",
            "陶", "姜", "戚", "谢", "邹", "喻", "柏", "水", "窦", "章",
            "云", "苏", "潘", "葛", "奚", "范", "彭", "郎", "鲁", "韦",
            "昌", "马", "苗", "凤", "花", "方", "俞", "任", "袁", "柳",
            "酆", "鲍", "史", "唐", "费", "廉", "岑", "薛", "雷", "贺",
            "倪", "汤", "滕", "殷", "罗", "毕", "郝", "邬", "安", "常",
            "乐", "于", "时", "傅", "皮", "卞", "齐", "康", "伍", "余",
            "元", "卜", "顾", "孟", "平", "黄", "和", "穆", "萧", "尹",
            "姚", "邵", "湛", "汪", "祁", "毛", "禹", "狄", "米", "贝",
            "明", "臧", "计", "伏", "成", "戴", "谈", "宋", "茅", "庞",
            "熊", "纪", "舒", "屈", "项", "祝", "董", "梁", "杜", "阮",
            "蓝", "闵", "席", "季", "麻", "强", "贾", "路", "娄", "危",
            "江", "童", "颜", "郭", "梅", "盛", "林", "刁", "钟", "徐",
            "邱", "骆", "高", "夏", "蔡", "田", "樊", "胡", "凌", "霍",
            "虞", "万", "支", "柯", "昝", "管", "卢", "莫", "经", "房",
            "裘", "缪", "干", "解", "应", "宗", "丁", "宣", "贲", "邓",
            "郁", "单", "杭", "洪", "包", "诸", "左", "石", "崔", "吉",
            "钮", "龚", "程", "嵇", "邢", "滑", "裴", "陆", "荣", "翁",
            "荀", "羊", "於", "惠", "甄", "曲", "家", "封", "芮", "羿",
            "储", "靳", "汲", "邴", "糜", "松", "井", "段", "富", "巫",
            "乌", "焦", "巴", "弓", "牧", "隗", "山", "谷", "车", "侯",
            "宓", "蓬", "全", "郗", "班", "仰", "秋", "仲", "伊", "宫",
            "宁", "仇", "栾", "暴", "甘", "钭", "厉", "戎", "祖", "武",
            "符", "刘", "景", "詹", "束", "龙", "叶", "幸", "司", "韶",
            "郜", "黎", "蓟", "薄", "印", "宿", "白", "怀", "蒲", "邰",
            "从", "鄂", "索", "咸", "籍", "赖", "卓", "蔺", "屠", "蒙",
            "池", "乔", "阴", "郁", "胥", "能", "苍", "双", "闻", "莘",
            "党", "翟", "谭", "贡", "劳", "逄", "姬", "申", "扶", "堵",
            "冉", "宰", "郦", "雍", "却", "璩", "桑", "桂", "濮", "牛",
            "寿", "通", "边", "扈", "燕", "冀", "郏", "浦", "尚", "农",
            "温", "别", "庄", "晏", "柴", "瞿", "阎", "充", "慕", "连",
            "茹", "习", "宦", "艾", "鱼", "容", "向", "古", "易", "慎",
            "戈", "廖", "庾", "终", "暨", "居", "衡", "步", "都", "耿",
            "满", "弘", "匡", "国", "文", "寇", "广", "禄", "阙", "东",
            "殴", "殳", "沃", "利", "蔚", "越", "夔", "隆", "师", "巩",
            "厍", "聂", "晁", "勾", "敖", "融", "冷", "訾", "辛", "阚",
            "那", "简", "饶", "空", "曾", "毋", "沙", "乜", "养", "鞠",
            "须", "丰", "巢", "关", "蒯", "相", "查", "后", "荆", "红",
            "游", "竺", "权", "逯", "盖", "益", "桓", "公", "万俟", "司马",
            "上官", "欧阳", "夏侯", "诸葛", "闻人", "东方", "赫连", "皇甫",
            "尉迟", "公羊", "澹台", "公冶", "宗政", "濮阳", "淳于", "单于",
            "太叔", "申屠", "公孙", "仲孙", "轩辕", "令狐", "钟离", "宇文",
            "长孙", "慕容", "鲜于", "闾丘", "司徒", "司空", "亓官", "司寇",
            "仉", "督", "子车", "颛孙", "端木", "巫马", "公西", "漆雕",
            "乐正", "壤驷", "公良", "拓跋", "夹谷", "宰父", "谷梁", "段干",
            "百里", "东郭", "南门", "呼延", "归", "海", "羊舌", "微生",
            "岳", "帅", "缑", "亢", "况", "后", "有", "琴", "梁丘", "左丘",
            "东门", "西门", "商", "牟", "佘", "佴", "伯", "赏", "南宫",
            "墨", "哈", "谯", "笪", "年", "爱", "阳", "佟"
        }
        return common_surnames
    
    def _load_sensitive_patterns(self) -> Dict[str, re.Pattern]:
        """加载敏感信息正则模式"""
        return {
            # 手机号（中国大陆）
            "phone": re.compile(r'1[3-9]\d{9}'),
            
            # 身份证号（18位）
            "id_card": re.compile(r'\d{17}[\dXx]'),
            
            # 银行卡号（16-19 位，支持空格或连字符分组）
            "bank_card": re.compile(r'(?<!\d)(?:\d[ -]?){15,18}\d(?!\d)'),
            
            # 邮箱地址
            "email": re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'),
            
            # 固定电话
            # 避免将统一社会信用代码、组织机构代码中的数字片段误判为固定电话。
            "landline": re.compile(r'(?<![A-Za-z0-9])(?:0\d{2,3}[-\s]?)?\d{7,8}(?![A-Za-z0-9])'),
            
            # IP 地址
            "ip_address": re.compile(r'\b(?:\d{1,3}\.){3}\d{1,3}\b'),
            # 允许 IPv6 的压缩写法（::），具体合法性由 ipaddress 校验。
            "ipv6_address": re.compile(r'(?<![0-9A-Fa-f:])(?:[0-9A-Fa-f]{0,4}:){2,7}[0-9A-Fa-f:]{0,4}(?![0-9A-Fa-f:])'),
            "mac_address": re.compile(r'\b(?:[0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2}\b'),
            
            # 车牌号
            "license_plate": re.compile(r'[京津沪渝冀豫云辽黑湘皖鲁新苏浙赣鄂桂甘晋蒙陕吉闽贵粤川青藏琼宁][A-Z][A-HJ-NP-Z0-9]{4,5}[A-HJ-NP-Z0-9挂学警港澳]'),
            
            # 护照号
            "passport": re.compile(r'\b(?:[EGPDST]\d{8}|[A-Z]{2}\d{7})\b'),
            "hong_kong_macao_permit": re.compile(r'\b(?:[CHM]\d{8,10}|[A-Z]{2}\d{6,8})\b'),
            
            # 军官证
            "military_id": re.compile(r'(?:军官证|军人证|军官身份号码)\s*[:：#]?\s*([A-Za-z0-9]{6,12})'),
            
            # 统一社会信用代码
            "unified_social_credit_code": re.compile(r'[0-9A-HJ-NPQRTUWXY]{2}\d{6}[0-9A-HJ-NPQRTUWXY]{10}'),
            "organization_code": re.compile(r'\b[0-9A-Z]{8}-?[0-9X]\b'),
            "business_license": re.compile(r'\b(?:\d{15}|[0-9A-HJ-NPQRTUWXY]{18})\b'),
            "jdbc_connection": re.compile(r'jdbc:[a-zA-Z0-9]+(?::[^\s"\']+)+', re.IGNORECASE),
            "date": re.compile(r'(?<!\d)(?:19|20)\d{2}[-/.年](?:0?[1-9]|1[0-2])[-/.月](?:0?[1-9]|[12]\d|3[01])(?:日)?(?!\d)'),
            "vehicle_identification_number": re.compile(r'\b[A-HJ-NPR-Z0-9]{17}\b'),
            "gender": re.compile(r'(?:性别|gender)\s*[:：]?\s*(?:男|女|未知)', re.IGNORECASE),
            "ethnicity": re.compile(r'(?:民族|族别)\s*[:：]?\s*[\u4e00-\u9fa5]{1,8}族?'),
            "province": re.compile(r'(?:省份|籍贯|所在省)\s*[:：]?\s*(?:北京市|天津市|上海市|重庆市|河北省|山西省|辽宁省|吉林省|黑龙江省|江苏省|浙江省|安徽省|福建省|江西省|山东省|河南省|湖北省|湖南省|广东省|海南省|四川省|贵州省|云南省|陕西省|甘肃省|青海省|台湾省|内蒙古自治区|广西壮族自治区|西藏自治区|宁夏回族自治区|新疆维吾尔自治区)'),
            "address": re.compile(r'(?:地址|住址|开户地址|收货地址)\s*[:：]?\s*[\u4e00-\u9fa5A-Za-z0-9#\-]{6,80}|(?<=为)[\u4e00-\u9fa5]{2,8}省[\u4e00-\u9fa5]{2,8}市[\u4e00-\u9fa5]{2,8}(?:区|县|镇|街道)[\u4e00-\u9fa5A-Za-z0-9#\-]{1,20}(?:路|街|巷|道)\d{1,5}号'),
            
            # 中文姓名（2-4个字，常见姓氏开头）
            "chinese_name": None,  # 需要特殊处理
            
            # 金额（万元、元等）
            "amount": re.compile(r'(?:\d{1,3}(?:,\d{3})+|\d+)(?:\.\d{1,2})?\s*(?:万元|元|美元|USD|CNY|￥|\$)'),
        }
    
    def extract_text(self, content: bytes, file_ext: str) -> str:
        """
        从不同格式的文件中提取文本
        
        Args:
            content: 文件内容（字节）
            file_ext: 文件扩展名
            
        Returns:
            提取的文本内容
        """
        try:
            if file_ext in ['.txt', '.csv', '.json', '.md']:
                # 文本文件直接解码
                return content.decode('utf-8', errors='ignore')
            
            elif file_ext == '.pdf':
                # PDF 文件处理
                return self._extract_text_from_pdf(content)
            
            elif file_ext == '.docx':
                # Word 文档处理
                return self._extract_text_from_docx(content)
            
            elif file_ext in ['.xlsx', '.xls']:
                # Excel 文件处理
                return self._extract_text_from_excel(content, file_ext)
            
            else:
                return ""
                
        except Exception as e:
            print(f"提取文本时出错：{str(e)}")
            return ""
    
    def _extract_text_from_pdf(self, content: bytes) -> str:
        """从 PDF 提取文本"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            print(f"PDF 解析错误：{str(e)}")
            return ""
    
    def _extract_text_from_docx(self, content: bytes) -> str:
        """从 Word 文档按原文档顺序提取段落与表格文本。"""
        try:
            doc = Document(io.BytesIO(content))
            blocks = []
            body = doc.element.body
            paragraph_by_element = {paragraph._p: paragraph for paragraph in doc.paragraphs}
            table_by_element = {table._tbl: table for table in doc.tables}

            for child in body.iterchildren():
                if child in paragraph_by_element:
                    blocks.append(paragraph_by_element[child].text)
                elif child in table_by_element:
                    table = table_by_element[child]
                    for row in table.rows:
                        blocks.append("\t".join(cell.text.replace("\n", " ") for cell in row.cells))

            return "\n".join(blocks).strip()
        except Exception as e:
            print(f"Word 文档解析错误：{str(e)}")
            return ""
    
    def _extract_text_from_excel(self, content: bytes, file_ext: str) -> str:
        """从 Excel 提取文本"""
        try:
            if file_ext != '.xlsx':
                raise ValueError('当前仅支持 XLSX 格式')
            workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            values = []
            for worksheet in workbook.worksheets:
                for row in worksheet.iter_rows(values_only=True):
                    values.extend(str(cell) for cell in row if cell is not None)
            return "\n".join(values).strip()
        except Exception as e:
            print(f"Excel 解析错误：{str(e)}")
            return ""
    
    def detect_sensitive_info(self, text: str, custom_rules: Optional[List[Dict[str, Any]]] = None) -> List[Dict[str, Any]]:
        """
        检测文本中的敏感信息
        
        Args:
            text: 要检测的文本
            
        Returns:
            检测到的敏感信息列表
        """
        detections = []
        
        # 1. 使用正则表达式检测
        regex_detections = self._detect_with_regex(text)
        detections.extend(regex_detections)
        
        # 2. 使用 Microsoft Presidio 检测（英文 PII）
        presidio_detections = self._detect_with_presidio(text)
        detections.extend(presidio_detections)
        
        # 3. 使用百家姓库检测中文姓名
        name_detections = self._detect_chinese_names(text)
        detections.extend(name_detections)

        # 4. 应用桌面客户端传入的本地规则（姓名、关键词或正则）
        detections.extend(self._detect_with_custom_rules(text, custom_rules or []))
        
        # 去重和合并
        detections = self._merge_detections(detections)
        
        # 更新统计信息
        self.stats["total_detections"] = len(detections)
        self.stats["by_type"] = {}
        for det in detections:
            det_type = det["type"]
            self.stats["by_type"][det_type] = self.stats["by_type"].get(det_type, 0) + 1
        
        return detections

    def _detect_with_custom_rules(self, text: str, rules: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """检测用户在本机规则面板中维护的敏感字段。"""
        detections = []
        for rule in rules[:100]:
            if not isinstance(rule, dict) or not rule.get("enabled", True):
                continue
            value = str(rule.get("value", "")).strip()
            if not value or len(value) > 500:
                continue
            rule_id = re.sub(r'[^A-Za-z0-9_]', '_', str(rule.get("id", "custom")))[:40] or "custom"
            try:
                pattern = re.compile(value if rule.get("kind") == "regex" else re.escape(value))
            except re.error:
                continue
            for match in pattern.finditer(text):
                if not match.group():
                    continue
                detections.append({
                    "type": rule_id,
                    "value": match.group(),
                    "start": match.start(),
                    "end": match.end(),
                    "confidence": 1.0,
                    "source": "local_rule",
                })
        return detections
    
    def _detect_with_regex(self, text: str) -> List[Dict[str, Any]]:
        """使用正则表达式检测敏感信息"""
        detections = []
        
        for pattern_name, pattern in self.sensitive_patterns.items():
            if pattern is None:
                continue
            
            matches = pattern.finditer(text)
            for match in matches:
                value = match.group(1) if pattern_name == "military_id" and match.lastindex else match.group()
                is_valid = self._is_valid_sensitive_value(pattern_name, value)
                # 企业证照测试数据有时使用虚构校验位；仍标记候选值，交由人工复核，
                # 避免用户框选后才得到“普通区域”而失去字段类型。
                if not is_valid and pattern_name not in {
                    "organization_code", "unified_social_credit_code", "business_license", "vehicle_identification_number"
                }:
                    continue
                detections.append({
                    "type": pattern_name,
                    "value": value,
                    "start": match.start(1) if pattern_name == "military_id" and match.lastindex else match.start(),
                    "end": match.end(1) if pattern_name == "military_id" and match.lastindex else match.end(),
                    "confidence": 0.9 if is_valid else 0.6,
                    "source": "regex"
                })
        
        return detections

    def _is_valid_sensitive_value(self, pattern_name: str, value: str) -> bool:
        """对需要校验的敏感标识执行算法校验，减少纯正则误报。"""
        validators = {
            "ip_address": self._is_valid_ipv4,
            "ipv6_address": self._is_valid_ipv6,
            "bank_card": self._is_valid_luhn,
            "id_card": self._is_valid_chinese_id,
            "organization_code": self._is_valid_organization_code,
            "unified_social_credit_code": self._is_valid_unified_credit_code,
            "business_license": self._is_valid_business_license,
            "vehicle_identification_number": self._is_valid_vin,
        }
        validator = validators.get(pattern_name)
        return validator(value) if validator else True

    @staticmethod
    def _is_valid_ipv4(value: str) -> bool:
        return all(0 <= int(part) <= 255 for part in value.split("."))

    @staticmethod
    def _is_valid_ipv6(value: str) -> bool:
        # 六组两位十六进制数是 MAC 地址，不作为 IPv6 处理。
        if re.fullmatch(r"(?:[0-9A-Fa-f]{2}:){5}[0-9A-Fa-f]{2}", value):
            return False
        try:
            ipaddress.IPv6Address(value)
            return True
        except ipaddress.AddressValueError:
            return False

    @staticmethod
    def _is_valid_luhn(value: str) -> bool:
        digits = re.sub(r"[ -]", "", value)
        if not digits.isdigit() or not 16 <= len(digits) <= 19:
            return False
        total = 0
        for index, char in enumerate(reversed(digits)):
            number = int(char)
            if index % 2:
                number *= 2
                if number > 9:
                    number -= 9
            total += number
        return total % 10 == 0

    @staticmethod
    def _is_valid_chinese_id(value: str) -> bool:
        value = value.upper()
        if not re.fullmatch(r"\d{17}[\dX]", value) or value[:2] == "00":
            return False
        try:
            datetime.strptime(value[6:14], "%Y%m%d")
        except ValueError:
            return False
        weights = (7, 9, 10, 5, 8, 4, 2, 1, 6, 3, 7, 9, 10, 5, 8, 4, 2)
        check_codes = "10X98765432"
        return check_codes[sum(int(char) * weight for char, weight in zip(value[:17], weights)) % 11] == value[-1]

    @staticmethod
    def _is_valid_organization_code(value: str) -> bool:
        value = value.replace("-", "")
        charset = "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        if len(value) != 9 or any(char not in charset for char in value):
            return False
        weights = (3, 7, 9, 10, 5, 8, 4, 2)
        remainder = sum(charset.index(char) * weight for char, weight in zip(value[:8], weights)) % 11
        expected = "X" if remainder == 10 else str(remainder)
        return expected == value[-1]

    @staticmethod
    def _is_valid_unified_credit_code(value: str) -> bool:
        charset = "0123456789ABCDEFGHJKLMNPQRTUWXY"
        value = value.upper()
        if len(value) != 18 or any(char not in charset for char in value):
            return False
        weights = (1, 3, 9, 27, 19, 26, 16, 17, 20, 29, 25, 13, 8, 24, 10, 30, 28)
        total = sum(charset.index(char) * weight for char, weight in zip(value[:17], weights))
        return charset[(31 - total % 31) % 31] == value[-1]

    def _is_valid_business_license(self, value: str) -> bool:
        return self._is_valid_unified_credit_code(value) if len(value) == 18 else self._is_valid_legacy_business_license(value)

    @staticmethod
    def _is_valid_legacy_business_license(value: str) -> bool:
        if not re.fullmatch(r"\d{15}", value):
            return False
        checksum = 10
        for char in value[:14]:
            checksum = (checksum + int(char)) % 10
            checksum = 10 if checksum == 0 else checksum
            checksum = (checksum * 2) % 11
        return (11 - checksum) % 10 == int(value[-1])

    @staticmethod
    def _is_valid_vin(value: str) -> bool:
        if len(value) != 17 or any(char in "IOQ" for char in value):
            return False
        values = {str(number): number for number in range(10)}
        values.update({
            "A": 1, "B": 2, "C": 3, "D": 4, "E": 5, "F": 6, "G": 7, "H": 8,
            "J": 1, "K": 2, "L": 3, "M": 4, "N": 5, "P": 7, "R": 9,
            "S": 2, "T": 3, "U": 4, "V": 5, "W": 6, "X": 7, "Y": 8, "Z": 9,
        })
        weights = (8, 7, 6, 5, 4, 3, 2, 10, 0, 9, 8, 7, 6, 5, 4, 3, 2)
        check = "0123456789X"
        return check[sum(values[char] * weight for char, weight in zip(value, weights)) % 11] == value[8]
    
    def _detect_with_presidio(self, text: str) -> List[Dict[str, Any]]:
        """使用 Microsoft Presidio 检测 PII"""
        detections = []
        
        if not self.analyzer:
            print("警告：Presidio 分析器不可用，跳过 Presidio 检测")
            return detections
        
        try:
            # 使用 Presidio 分析文本
            results = self.analyzer.analyze(
                text=text,
                language="en",
                entities=[
                    "PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER",
                    "US_SSN", "CREDIT_CARD", "IP_ADDRESS",
                    "LOCATION", "DATE_TIME"
                ]
            )
            
            for result in results:
                # 映射 Presidio 实体类型到我们的类型
                type_mapping = {
                    "PERSON": "name",
                    "EMAIL_ADDRESS": "email",
                    "PHONE_NUMBER": "phone",
                    "US_SSN": "ssn",
                    "CREDIT_CARD": "bank_card",
                    "IP_ADDRESS": "ip_address",
                    "LOCATION": "address",
                    "DATE_TIME": "date"
                }
                
                detection_type = type_mapping.get(result.entity_type, "other")
                
                detections.append({
                    "type": detection_type,
                    "value": text[result.start:result.end],
                    "start": result.start,
                    "end": result.end,
                    "confidence": result.score,
                    "source": "presidio"
                })
                
        except Exception as e:
            print(f"Presidio 检测错误：{str(e)}")
        
        return detections
    
    def _detect_chinese_names(self, text: str) -> List[Dict[str, Any]]:
        """使用百家姓库检测中文姓名（带上下文验证）"""
        detections = []
        
        # 常见的非人名词汇（以姓氏开头但不是人名）
        common_words = {
            # 金融商务
            "金融", "任何", "商业", "公司", "基金", "投资", "银行", "证券",
            "保险", "信托", "资产", "管理", "服务", "咨询", "集团", "有限",
            "股份", "企业", "机构", "中心", "发展", "建设", "科技", "技术",
            "信息", "网络", "文化", "教育", "医疗", "健康", "环保", "能源",
            
            # 行业领域
            "房地产", "物业", "酒店", "餐饮", "零售", "批发", "贸易", "进出口",
            "物流", "运输", "仓储", "制造", "生产", "加工", "研发", "设计",
            "工程", "建筑", "装饰", "材料", "设备", "机械", "电子", "电器",
            "汽车", "交通", "通信", "传媒", "广告", "娱乐", "旅游", "农业",
            "畜牧", "渔业", "林业", "矿业", "石油", "化工", "钢铁", "有色金属",
            "纺织", "服装", "食品", "饮料", "医药", "生物", "制药", "医院",
            
            # 教育机构
            "学校", "大学", "学院", "培训", "职业", "人才", "人力", "资源",
            
            # 专业服务
            "财务", "会计", "审计", "税务", "法律", "事务所", "律师", "公证",
            "评估", "评级", "担保", "典当", "拍卖", "租赁", "抵押", "信贷",
            
            # 金融市场
            "贷款", "融资", "债券", "股票", "期货", "期权", "外汇", "黄金",
            "白银", "铜", "铝", "锌", "镍", "锡", "铅", "钢材", "煤炭",
            "石油", "天然气", "电力", "水务", "燃气", "热力", "新能源", "太阳能",
            "风能", "水电", "核电", "生物质", "地热", "海洋", "潮汐", "储能",
            
            # 科技领域
            "电池", "充电桩", "电动汽车", "新能源汽车", "智能", "人工智能", "大数据",
            "云计算", "物联网", "区块链", "数字货币", "互联网", "电商", "电子商务",
            
            # 文娱传媒
            "社交", "游戏", "视频", "音频", "直播", "短视频", "长视频", "音乐",
            "电影", "电视", "动漫", "动画", "漫画", "小说", "文学", "艺术",
            "绘画", "雕塑", "摄影", "设计", "创意", "策划", "营销", "推广",
            "品牌", "形象", "包装", "展览", "展示", "会议", "活动", "赛事",
            
            # 体育健身
            "体育", "运动", "健身", "瑜伽", "舞蹈", "武术", "太极", "气功",
            
            # 生活服务
            "养生", "保健", "营养", "饮食", "烹饪", "厨艺", "美食", "料理",
            "咖啡", "茶", "酒", "饮料", "零食", "糖果", "巧克力", "饼干",
            "面包", "蛋糕", "甜点", "冰淇淋", "冷饮", "热饮", "果汁", "牛奶",
            "酸奶", "豆浆", "豆腐", "蔬菜", "水果", "肉类", "海鲜", "水产",
            
            # 农业食品
            "粮食", "谷物", "面粉", "大米", "小麦", "玉米", "大豆", "花生",
            "芝麻", "菜籽", "橄榄", "棕榈", "椰子", "可可", "咖啡豆", "茶叶",
            "烟草", "棉花", "羊毛", "丝绸", "皮革", "橡胶", "塑料", "玻璃",
            "陶瓷", "水泥", "木材", "竹材", "石材", "砂石", "砖瓦", "石灰",
            "石膏", "沥青", "涂料", "油漆", "染料", "颜料", "墨水", "油墨",
            "胶水", "粘合剂", "密封剂", "清洁剂", "消毒剂", "杀虫剂", "除草剂",
            "化肥", "农药", "种子", "种苗", "花卉", "苗木", "草坪", "园艺",
            
            # 动物相关
            "宠物", "狗", "猫", "鸟", "鱼", "虫", "兽", "禽", "畜",
            "野生动物", "保护", "救助", "领养", "寄养", "托管", "幼儿园", "早教",
            
            # 学历学位
            "小学", "中学", "高中", "大学", "研究生", "博士", "硕士", "学士",
            "学位", "学历", "文凭", "证书", "资格", "执照", "许可", "批准",
            
            # 行政管理
            "备案", "登记", "注册", "认证", "认可", "鉴定", "检验", "检测",
            "测试", "试验", "实验", "研究", "开发", "创新", "发明", "专利",
            "商标", "版权", "著作权", "知识产权", "商业秘密", "保密", "竞业",
            "禁止", "限制", "约束", "义务", "责任", "权利", "权益", "利益",
            
            # 财务会计
            "利润", "收益", "回报", "分红", "利息", "本金", "资本", "资金",
            "现金", "货币", "外汇", "汇率", "利率", "税率", "费率", "价格",
            "成本", "费用", "开支", "支出", "收入", "营收", "销售额", "营业额",
            "毛利", "净利", "税前", "税后", "扣除", "抵扣", "减免",
            
            # 商业交易
            "优惠", "折扣", "促销", "特价", "清仓", "甩卖", "拍卖", "竞拍",
            "招标", "投标", "中标", "落标", "废标", "流标", "围标", "串标",
            "陪标", "抢标", "夺标", "争标", "竞标", "比价", "议价", "砍价",
            "压价", "抬价", "涨价", "降价", "调价", "定价", "报价", "出价",
            "还价", "底价", "起价", "基价", "单价", "总价", "均价", "差价",
            
            # 物理单位
            "价差", "利差", "息差", "汇差", "时差", "温差", "压差", "电压",
            "电流", "电阻", "电容", "电感", "电荷", "电场", "磁场", "电磁",
            "辐射", "波长", "频率", "振幅", "相位", "周期", "速度", "加速度",
            "力", "质量", "密度", "体积", "面积", "长度", "宽度", "高度",
            "深度", "厚度", "直径", "半径", "周长", "圆周率", "角度", "弧度",
            "梯度", "度", "分", "秒", "毫秒", "微秒", "纳秒", "皮秒",
            "飞秒", "阿秒", "年", "月", "日", "时", "分", "秒", "周",
            "旬", "季度", "半年", "全年", "年度", "月份", "日期", "时刻",
            "时间", "时期", "阶段", "期间", "期限",
            
            # 合同法律术语
            "盖章", "签字", "签名", "签署", "签订", "订立", "生效", "失效",
            "终止", "解除", "撤销", "变更", "补充", "修改", "修订", "废止",
            "甲方", "乙方", "丙方", "丁方", "各方", "双方", "当事人", "当事人",
            "委托", "受托", "代理", "代表", "授权", "委托书", "授权书", "证明",
            "声明", "承诺", "保证", "担保", "抵押", "质押", "留置", "定金",
            "违约", "赔偿", "补偿", "罚款", "罚金", "滞纳金", "利息", "利率",
            "本金", "本金", "本金", "本金", "本金", "本金", "本金", "本金",
            "债务", "债权", "债务人", "债权人", "债务", "债权", "债务人", "债权人",
            "合同", "协议", "契约", "合约", "约定", "条款", "款项", "条件",
            "期限", "有效期", "生效日期", "失效日期", "签订日期", "签署日期",
            "地点", "签订地点", "签署地点", "履行地", "履行地点",
            "争议", "纠纷", "仲裁", "诉讼", "法院", "管辖", "适用法律",
            "不可抗力", "免责", "免责条款", "保密条款", "竞业限制", "竞业禁止",
            "知识产权", "著作权", "专利权", "商标权", "商业秘密", "技术秘密",
            "通知", "送达", "地址", "联系方式", "电话", "邮箱", "传真",
            "附件", "附录", "补充协议", "备忘录", "意向书", "框架协议",
            
            # 常见动词形容词
            "可以", "不能", "不得", "必须", "应当", "应该", "需要", "允许",
            "禁止", "限制", "约束", "规定", "约定", "说明", "解释", "定义",
            "包括", "包含", "涉及", "关于", "对于", "鉴于", "由于", "因为",
            "所以", "因此", "但是", "然而", "不过", "可是", "而是", "或者",
            "以及", "并且", "而且", "同时", "另外", "此外", "除此之外",
            "总之", "综上所述", "如上所述", "如下", "以上", "以下", "前述",
            "后述", "左述", "右述", "上述", "下述", "前述", "后述",
            
            # 时间日期
            "今天", "昨天", "明天", "后天", "前天", "本周", "上周", "下周",
            "本月", "上月", "下月", "本年", "去年", "明年", "年初", "年末",
            "月底", "月初", "月中", "周末", "工作日", "节假日", "休息日",
            
            # 常见名词
            "问题", "情况", "条件", "方式", "方法", "措施", "手段", "途径",
            "目标", "目的", "结果", "效果", "影响", "作用", "意义", "价值",
            "原则", "规则", "制度", "体制", "机制", "体系", "系统", "程序",
            "流程", "步骤", "阶段", "环节", "部分", "方面", "角度", "层面",
            "范围", "领域", "区域", "地区", "位置", "地点", "场所", "场合",
            "环境", "背景", "前提", "基础", "条件", "因素", "要素", "成分",
            "内容", "形式", "格式", "样式", "类型", "种类", "类别", "品种",
            "质量", "数量", "程度", "水平", "标准", "规格", "指标", "参数",
            "数据", "信息", "资料", "材料", "资源", "能源", "动力", "资金",
            "财产", "资产", "资本", "财富", "财富", "财富", "财富", "财富",
            
            # 金融投资常用词
            "平衡", "稳健", "进取", "保守", "激进", "成长", "价值", "收益",
            "风险", "回报", "波动", "分散", "集中", "配置", "组合", "策略",
            "分析", "研判", "预测", "趋势", "行情", "走势", "涨跌", "反弹",
            "回调", "突破", "支撑", "压力", "均线", "指标", "信号", "形态",
            "基本面", "技术面", "政策面", "资金面", "消息面", "市场面",
            "利好", "利空", "利多", "看涨", "看跌", "看多", "看空", "做多",
            "做空", "持仓", "建仓", "加仓", "减仓", "平仓", "清仓", "调仓",
            "满仓", "半仓", "空仓", "轻仓", "重仓", "底仓", "浮仓", "锁仓",
            "短线", "中线", "长线", "波段", "日内", "趋势", "震荡", "盘整",
            "突破", "跌破", "站上", "失守", "守住", "承压", "受阻", "企稳",
            "回升", "反弹", "回落", "下跌", "上涨", "横盘", "整理", "蓄势",
            "放量", "缩量", "天量", "地量", "温和", "急剧", "缓慢", "快速",
            "涨停", "跌停", "熔断", "停牌", "复牌", "摘牌", "退市", "上市",
            "发行", "申购", "配售", "中签", "缴款", "上市", "流通", "限售",
            "解禁", "减持", "增持", "回购", "分红", "送股", "转增", "配股",
            "增发", "可转债", "优先股", "普通股", "A股", "B股", "H股", "美股",
            "港股", "沪深", "上证", "深证", "创业板", "科创板", "北交所", "新三板",
            "主板", "中小板", "指数", "成分", "权重", "样本", "调整", "纳入",
            "剔除", "扩容", "缩容", "活跃", "低迷", "火爆", "冷清", "清淡",
            "拥挤", "分化", "轮动", "切换", "风格", "偏好", "情绪", "信心",
            "恐慌", "贪婪", "乐观", "悲观", "中性", "谨慎", "积极", "消极",
            "主动", "被动", "量化", "对冲", "套利", "投机", "投资", "理财",
            "储蓄", "存款", "贷款", "融资", "融券", "杠杆", "保证金", "抵押",
            "质押", "担保", "信用", "评级", "评级", "评级", "评级", "评级",
        }
        
        # 上下文关键词（姓名前后出现的词）
        # 仅保留明确的人名语境；泛化连接词和括号说明会把“支付”“验证”等普通词误判为姓名。
        context_before = ["姓名", "联系人", "项目联系人", "负责人", "申请人", "经办人", "收件人"]
        context_after = ["先生", "女士", "小姐", "同志", "同学", "老师", "教授", "医生", "律师", "工程师", "经理", "总监", "总裁", "董事长", "主任", "处长", "局长", "部长", "省长", "市长", "县长", "镇长", "村长", "书记", "主席", "总理", "首相", "总统", "国王", "女王", "皇帝", "皇后", "王子", "公主", "贵族", "爵士", "骑士", "将军", "元帅", "上将", "中将", "少将", "大校", "上校", "中校", "少校", "上尉", "中尉", "少尉", "军士长", "上士", "中士", "下士", "列兵", "新兵", "学员", "士兵", "军人", "军官", "将领", "统帅", "指挥官", "司令官", "参谋长", "政委", "指导员", "教导员", "连长", "排长", "班长", "组长", "队长", "大队长", "中队长", "小队长", "支队长", "总队长", "师长", "旅长", "团长", "营长", "连长", "排长", "班长", "组长", "队长", "大队长", "中队长", "小队长", "支队长", "总队长", "师长", "旅长", "团长", "营长"]
        
        # 使用 jieba 分词
        words = list(jieba.cut(text))
        
        current_pos = 0
        for i, word in enumerate(words):
            # 查找在原文中的位置
            start = text.find(word, current_pos)
            if start == -1:
                continue
            
            # 检查是否是中文姓名（2-4个字，以常见姓氏开头）
            if (len(word) >= 2 and len(word) <= 4 and 
                word[0] in self.chinese_surnames and
                all('\u4e00' <= c <= '\u9fff' for c in word) and
                word not in common_words):
                
                # 检查上下文
                is_name = False
                
                # 检查前面的词
                if i > 0:
                    prev_word = words[i-1]
                    if prev_word in context_before:
                        is_name = True
                
                # 检查后面的词
                if i < len(words) - 1:
                    next_word = words[i+1]
                    if next_word in context_after:
                        is_name = True
                
                # 如果有上下文支持，则认为是姓名
                if is_name:
                    detections.append({
                        "type": "chinese_name",
                        "value": word,
                        "start": start,
                        "end": start + len(word),
                        "confidence": 0.8,
                        "source": "chinese_surnames"
                    })
            
            current_pos = start + len(word)
        
        return detections
    
    def _merge_detections(self, detections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """合并重叠的检测结果"""
        if not detections:
            return []
        
        # 结构化标识优先于启发式姓名、固定电话等宽泛规则。
        priority = {
            "ipv6_address": 100, "ip_address": 100, "mac_address": 100,
            "bank_card": 98, "id_card": 98, "unified_social_credit_code": 97,
            "organization_code": 97, "business_license": 97,
            "vehicle_identification_number": 96, "passport": 99,
            "hong_kong_macao_permit": 99, "military_id": 99,
            "email": 94, "jdbc_connection": 94, "address": 90,
            "phone": 88, "landline": 50, "chinese_name": 20,
        }
        # 按位置排序；同一起点先处理优先级高、范围完整的候选。
        detections.sort(key=lambda x: (x["start"], -priority.get(x["type"], 60), -x["end"]))
        
        merged = []
        current = detections[0]
        
        for next_det in detections[1:]:
            # 检查是否重叠
            if next_det["start"] < current["end"]:
                # 优先保留结构化规则；同优先级再比较置信度与覆盖范围。
                current_priority = priority.get(current["type"], 60)
                next_priority = priority.get(next_det["type"], 60)
                if (next_priority, next_det["confidence"], next_det["end"] - next_det["start"]) > (current_priority, current["confidence"], current["end"] - current["start"]):
                    current = next_det
            else:
                merged.append(current)
                current = next_det
        
        merged.append(current)
        
        return merged
    
    def redact_text(self, text: str, custom_rules: Optional[List[Dict[str, Any]]] = None) -> Dict[str, Any]:
        """
        执行文本脱敏
        
        Args:
            text: 要脱敏的文本
            custom_rules: 本机敏感字段规则
            
        Returns:
            脱敏结果，包含脱敏文本和映射表
        """
        # 检测敏感信息
        detections = self.detect_sensitive_info(text, custom_rules)
        
        # 创建映射表
        mappings = []
        redacted_text = text
        
        # 按位置从后往前替换，避免位置偏移
        for i, detection in enumerate(reversed(detections)):
            # 使用 {type_001} 格式，更容易识别和匹配
            det_type = detection['type'].upper()
            placeholder = f"{{{det_type}_{i+1:03d}}}"
            
            # 记录映射关系
            mappings.append({
                "id": i + 1,
                "placeholder": placeholder,
                "type": detection["type"],
                "original": detection["value"],
                "start": detection["start"],
                "end": detection["end"],
                "confidence": detection["confidence"],
                "source": detection["source"]
            })
            
            # 执行替换
            redacted_text = (
                redacted_text[:detection["start"]] + 
                placeholder + 
                redacted_text[detection["end"]:]
            )
        
        # 反转映射表（按位置从前到后）
        mappings.reverse()
        
        return {
            "success": True,
            "original_text": text,
            "redacted_text": redacted_text,
            "original_length": len(text),
            "redacted_length": len(redacted_text),
            "detection_count": len(detections),
            "mappings": mappings,
            "created_at": datetime.now().isoformat(),
            "stats": self.stats
        }
    
    def restore_text(self, redacted_text: str, mapping_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        还原脱敏文本
        
        Args:
            redacted_text: 脱敏后的文本
            mapping_data: 映射表数据
            
        Returns:
            还原结果
        """
        import re
        
        mappings = mapping_data.get("mappings", [])
        
        restored_text = redacted_text
        
        # 按占位符长度从长到短排序，避免部分匹配
        sorted_mappings = sorted(mappings, key=lambda x: len(x.get("placeholder", "")), reverse=True)
        
        # 记录替换统计
        replacements = 0
        failed = []
        
        for mapping in sorted_mappings:
            placeholder = mapping.get("placeholder", "")
            original = mapping.get("original", "")
            
            if not placeholder or original is None:
                continue
            
            # 使用正则表达式进行精确匹配
            # 转义花括号 { 和 }
            escaped_placeholder = re.escape(placeholder)
            
            # 检查是否存在于文本中
            if re.search(escaped_placeholder, restored_text):
                restored_text = re.sub(escaped_placeholder, original, restored_text)
                replacements += 1
            else:
                # 记录未找到的占位符
                failed.append(placeholder)
        
        # 清理可能残留的占位符格式 {TYPE_数字}
        # 匹配形如 {大写字母_数字} 的残留
        restored_text = re.sub(r'\{[A-Z]+_\d{3}\}', '', restored_text)
        
        # 清理可能残留的单个大括号（前后没有内容的情况）
        # 例如：}内容{ 或 {内容 或 内容}
        restored_text = re.sub(r'\}(\s*)\{', r'\1', restored_text)
        restored_text = re.sub(r'^\}', '', restored_text)
        restored_text = re.sub(r'\{$', '', restored_text)
        
        # 清理连续的大括号
        restored_text = re.sub(r'\{\s*\}', '', restored_text)
        
        # 清理可能残留的不完整占位符（如 {TYPE_ 没有闭合）
        restored_text = re.sub(r'\{[A-Z_]+(?=_\d{3}\})', '', restored_text)
        
        # 清理可能残留的数字片段（如 _018}）
        restored_text = re.sub(r'_\d{3}\}', '', restored_text)
        
        return {
            "success": True,
            "restored_text": restored_text,
            "redacted_length": len(redacted_text),
            "restored_length": len(restored_text),
            "mappings_applied": replacements,
            "mappings_failed": failed,
            "created_at": datetime.now().isoformat()
        }
