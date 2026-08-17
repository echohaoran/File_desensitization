"""
多步骤脱敏系统 - Python 后端服务
使用 FastAPI 提供文件脱敏 API
集成 Microsoft Presidio 和中国百家姓库进行 PII 检测
"""

import os
import json
import io
import uuid
import tempfile
import subprocess
import shutil
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from starlette.background import BackgroundTask

from desensitization_service import DesensitizationService

app = FastAPI(
    title="多步骤脱敏系统 API",
    description="面向企业内部投资人员与私募基金管理员的敏感数据脱敏工具",
    version="1.0.0"
)

# 配置 CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 生产环境应限制为特定域名
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 初始化脱敏服务
desensitization_service = DesensitizationService()

# 用户上传文件存储目录
APP_ROOT = Path(os.getenv("DESENS_APP_ROOT", Path(__file__).resolve().parent.parent))
UPLOAD_DIR = os.getenv("UPLOAD_DIR", str(APP_ROOT / "uploads"))
os.makedirs(UPLOAD_DIR, exist_ok=True)
UPDATE_SOURCES = {
    "github": {
        "name": "GitHub",
        "repository": os.getenv("GITHUB_UPDATE_REPOSITORY", "https://github.com/echohaoran/File_desensitization.git"),
    },
    "gitee": {
        "name": "Gitee",
        "repository": os.getenv("GITEE_UPDATE_REPOSITORY", "https://gitee.com/echohaoran/file_desensitization.git"),
    },
    "cnb": {
        "name": "CNB",
        "repository": os.getenv("CNB_UPDATE_REPOSITORY", "https://cnb.cool/echohaoran/File_desensitization.git"),
    },
}

REDACTION_NOTICE = (
    "【处理与还原规则】本文件包含脱敏占位符（例如 {PHONE_001}）。"
    "请用户及任何 AI/Agent 在计算、分析、编辑、改写或排版时完整保留占位符的花括号、字段名与编号，不得删除、拆分、改写或替换；否则可能无法还原原始内容。"
)


def _add_docx_redaction_notice(document) -> None:
    """在 Word 文档正文顶部插入可供用户和 Agent 阅读的还原规则。"""
    paragraph = document.paragraphs[0].insert_paragraph_before() if document.paragraphs else document.add_paragraph()
    paragraph.add_run(REDACTION_NOTICE).bold = True
    paragraph.paragraph_format.space_after = 12


def _add_xlsx_redaction_notice(workbook) -> None:
    """在每个工作表顶端插入脱敏占位符保留规则。"""
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    for worksheet in workbook.worksheets:
        worksheet.insert_rows(1)
        end_column = max(1, worksheet.max_column)
        worksheet.cell(1, 1, REDACTION_NOTICE)
        worksheet.cell(1, 1).font = Font(bold=True, color="9C2F2F")
        worksheet.cell(1, 1).fill = PatternFill("solid", fgColor="FFF3CD")
        worksheet.cell(1, 1).alignment = Alignment(wrap_text=True, vertical="center")
        worksheet.row_dimensions[1].height = 48
        if end_column > 1:
            worksheet.merge_cells(f"A1:{get_column_letter(end_column)}1")


def parse_custom_rules(raw_rules: str) -> List[Dict[str, Any]]:
    """解析客户端传入的本机规则；无效内容按空规则处理。"""
    try:
        rules = json.loads(raw_rules or "[]")
        return rules if isinstance(rules, list) else []
    except json.JSONDecodeError:
        return []


def _remove_directory(path: str) -> None:
    """在响应传输完成后清理单次转换产生的临时目录。"""
    shutil.rmtree(path, ignore_errors=True)


def _safe_upload_name(filename: Optional[str], fallback: str) -> str:
    """只保留文件名，避免上传名称影响临时目录中的路径。"""
    return Path(filename or fallback).name


def build_docx_preview(content: bytes) -> List[Dict[str, Any]]:
    """返回供前端渲染的 Word 文档结构，并保留与抽取文本一致的字符偏移。"""
    from docx import Document

    doc = Document(io.BytesIO(content))
    blocks: List[Dict[str, Any]] = []
    cursor = 0
    paragraph_by_element = {paragraph._p: paragraph for paragraph in doc.paragraphs}
    table_by_element = {table._tbl: table for table in doc.tables}

    def add_text(value: str) -> Dict[str, Any]:
        nonlocal cursor
        item = {"text": value, "start": cursor, "end": cursor + len(value)}
        cursor += len(value) + 1
        return item

    for child in doc.element.body.iterchildren():
        if child in paragraph_by_element:
            paragraph = paragraph_by_element[child]
            style_name = (paragraph.style.name or '').lower() if paragraph.style else ''
            font_sizes = [run.font.size.pt for run in paragraph.runs if run.font.size]
            largest_font_size = max(font_sizes, default=None)
            is_bold = any(bool(run.bold) for run in paragraph.runs)
            kind = 'heading' if (
                style_name.startswith('heading') or '标题' in style_name or style_name == 'title'
                or (is_bold and largest_font_size and largest_font_size >= 16)
            ) else 'paragraph'
            level = next((int(char) for char in style_name if char.isdigit()), 0)
            alignment = {1: 'center', 2: 'right', 3: 'justify'}.get(paragraph.alignment)
            is_list = paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
            blocks.append({
                "type": kind, "level": level,
                "format": {
                    "font_size": largest_font_size,
                    "bold": is_bold,
                    "alignment": alignment,
                    "list": is_list,
                },
                **add_text(paragraph.text),
            })
        elif child in table_by_element:
            rows = []
            for row in table_by_element[child].rows:
                rows.append([add_text(cell.text.replace("\n", " ")) for cell in row.cells])
            blocks.append({"type": "table", "rows": rows})

    return blocks


def _iter_docx_text_targets(document):
    """按正文元素顺序返回可替换的段落和表格单元格。"""
    paragraph_by_element = {paragraph._p: paragraph for paragraph in document.paragraphs}
    table_by_element = {table._tbl: table for table in document.tables}
    for child in document.element.body.iterchildren():
        if child in paragraph_by_element:
            yield paragraph_by_element[child]
        elif child in table_by_element:
            for row in table_by_element[child].rows:
                for cell in row.cells:
                    yield cell


def _redact_fragment(text: str, offset: int, mappings: List[Dict[str, Any]]) -> str:
    """按原文全局字符偏移替换一个段落或单元格，避免重复值错配。"""
    end = offset + len(text)
    targets = sorted(
        (item for item in mappings if item.get('start', -1) >= offset and item.get('end', -1) <= end),
        key=lambda item: item['start'], reverse=True,
    )
    redacted = text
    for item in targets:
        start = item['start'] - offset
        stop = item['end'] - offset
        if redacted[start:stop] == item.get('original', ''):
            redacted = redacted[:start] + item['placeholder'] + redacted[stop:]
    return redacted


def _replace_target_text(target, value: str) -> None:
    """替换文字同时保留段落/单元格/表格本身。"""
    if hasattr(target, 'paragraphs'):
        paragraph = target.paragraphs[0] if target.paragraphs else target.add_paragraph()
        for extra in target.paragraphs[1:]:
            extra.clear()
    else:
        paragraph = target
    paragraph.clear()
    paragraph.add_run(value)


def redact_file_preserving_format(content: bytes, file_ext: str, mappings: List[Dict[str, Any]]) -> tuple[bytes, str]:
    """输出不丢失段落、表格和工作表的脱敏文档；PDF 输出 DOCX。"""
    if file_ext == '.pdf':
        from pdf2docx import Converter
        with tempfile.TemporaryDirectory(prefix='pdf_redact_') as temp_dir:
            pdf_path = os.path.join(temp_dir, 'source.pdf')
            docx_path = os.path.join(temp_dir, 'source.docx')
            Path(pdf_path).write_bytes(content)
            converter = Converter(pdf_path)
            try:
                converter.convert(docx_path)
            finally:
                converter.close()
            content = Path(docx_path).read_bytes()
        file_ext = '.docx'

    if file_ext == '.docx':
        from docx import Document
        document = Document(io.BytesIO(content))
        _add_docx_redaction_notice(document)
        offset = 0
        targets = list(_iter_docx_text_targets(document))
        applied_ids = set()
        for target in targets:
            original = target.text.replace('\n', ' ') if hasattr(target, 'paragraphs') else target.text
            scoped_mappings = [
                item for item in mappings
                if item.get('start', -1) >= offset and item.get('end', -1) <= offset + len(original)
            ]
            replacement = _redact_fragment(original, offset, scoped_mappings)
            if replacement != original:
                _replace_target_text(target, replacement)
                applied_ids.update(item.get('id') for item in scoped_mappings)
            offset += len(original) + 1

        # PDF 每次转换的字符切分可能不同。保留精确偏移优先级，并对未命中的项按原值兜底。
        for item in mappings:
            if item.get('id') in applied_ids:
                continue
            for target in targets:
                original = target.text.replace('\n', ' ') if hasattr(target, 'paragraphs') else target.text
                if item.get('original') and item['original'] in original:
                    _replace_target_text(target, original.replace(item['original'], item['placeholder'], 1))
                    break
        output = io.BytesIO()
        document.save(output)
        return output.getvalue(), '.docx'

    if file_ext == '.xlsx':
        from openpyxl import load_workbook
        workbook = load_workbook(io.BytesIO(content))
        offset = 0
        for worksheet in workbook.worksheets:
            for row in worksheet.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    original = str(cell.value)
                    replacement = _redact_fragment(original, offset, mappings)
                    if replacement != original:
                        cell.value = replacement
                    offset += len(original) + 1
        _add_xlsx_redaction_notice(workbook)
        output = io.BytesIO()
        workbook.save(output)
        return output.getvalue(), '.xlsx'

    raise HTTPException(status_code=400, detail=f'暂不支持保留格式输出：{file_ext}')


def _find_soffice() -> Optional[str]:
    """查找 LibreOffice；桌面端将其作为 DOCX→PDF 的可选外部组件。"""
    configured = os.getenv("SOFFICE_PATH")
    candidates = [configured] if configured else []
    candidates.extend([
        shutil.which("soffice"),
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        r"C:\\Program Files\\LibreOffice\\program\\soffice.exe",
        r"C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
    ])
    return next((candidate for candidate in candidates if candidate and Path(candidate).exists()), None)


def _run_soffice(source_path: str, output_dir: str, output_format: str) -> subprocess.CompletedProcess:
    """使用独立 LibreOffice 用户配置执行一次转换，避免并发请求互相占用配置锁。"""
    soffice = _find_soffice()
    if not soffice:
        raise HTTPException(
            status_code=503,
            detail="Word 转 PDF 需要 LibreOffice。请安装 LibreOffice，或通过 SOFFICE_PATH 指定 soffice 可执行文件后重试。",
        )
    profile_dir = tempfile.mkdtemp(prefix="soffice_profile_")
    try:
        return subprocess.run(
            [
                soffice,
                f'-env:UserInstallation={Path(profile_dir).as_uri()}',
                '--headless',
                '--convert-to', output_format,
                '--outdir', output_dir,
                source_path,
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )
    finally:
        _remove_directory(profile_dir)


@app.get("/")
async def root():
    """API 根路径"""
    return {
        "message": "多步骤脱敏系统 API",
        "version": "1.0.0",
        "endpoints": {
            "/api/redact": "POST - 上传文件并执行初步脱敏",
            "/api/detect": "POST - 仅检测敏感信息（不脱敏）",
            "/api/health": "GET - 健康检查"
        }
    }


@app.get("/api/health")
async def health_check():
    """健康检查接口"""
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}


@app.get("/api/runtime/capabilities")
async def runtime_capabilities():
    """返回运行环境能力，供前端提前提示可选外部组件。"""
    soffice = _find_soffice()
    return {
        "word_to_pdf": {
            "available": bool(soffice),
            "requirement": "LibreOffice（可选；仅 DOCX→PDF 必需）",
        }
    }


@app.get("/api/version/check")
async def check_update(source: str = "github"):
    """从预设 GitHub、Gitee 或 CNB main 分支检查更新。"""
    source_key = source.lower().strip()
    source_config = UPDATE_SOURCES.get(source_key)
    if not source_config:
        raise HTTPException(status_code=400, detail="不支持的更新来源")
    project_root = APP_ROOT
    repository = source_config["repository"]
    try:
        try:
            current = subprocess.run(
                ["git", "-C", str(project_root), "rev-parse", "HEAD"],
                capture_output=True, text=True, timeout=5, check=True,
            ).stdout.strip()
        except subprocess.CalledProcessError:
            # 开发服务器可能通过文件同步部署而不携带 .git；使用随包发布的修订清单。
            manifest = json.loads((project_root / "public" / "version.json").read_text(encoding="utf-8"))
            current = str(manifest.get("revision") or "")
        if not current:
            raise RuntimeError("当前部署缺少 revision 信息")
        remote = subprocess.run(
            ["git", "ls-remote", repository, "refs/heads/main"],
            capture_output=True, text=True, timeout=12, check=True,
        ).stdout.strip().split()
        latest = remote[0] if remote else ""
        if not latest:
            raise RuntimeError(f"{source_config['name']} main 分支未返回提交记录")
        commits = []
        try:
            with tempfile.TemporaryDirectory(prefix=f"{source_key}_update_") as temp_dir:
                subprocess.run(
                    ["git", "clone", "--quiet", "--depth", "5", "--branch", "main", repository, temp_dir],
                    capture_output=True, text=True, timeout=25, check=True,
                )
                log = subprocess.run(
                    ["git", "-C", temp_dir, "log", "-5", "--pretty=format:%H%x1f%h%x1f%s%x1f%aI"],
                    capture_output=True, text=True, timeout=5, check=True,
                ).stdout
                commits = [
                    {"id": parts[0], "short_id": parts[1], "message": parts[2], "created_at": parts[3]}
                    for line in log.splitlines()
                    if len(parts := line.split("\x1f")) == 4
                ]
        except Exception:
            # 提交摘要不可用不影响更新状态判断。
            commits = []
        return {
            "source": source_config["name"],
            "source_key": source_key,
            "repository": repository.removesuffix(".git"),
            "branch": "main",
            "current_revision": current,
            "latest_revision": latest,
            "update_available": latest != current,
            "commits": commits,
        }
    except Exception as error:
        return JSONResponse(
            status_code=503,
            content={"source": "CNB", "detail": f"CNB 更新检查失败：{error}"},
        )


@app.post("/api/pdf-to-word")
async def convert_pdf_to_word(file: UploadFile = File(...)):
    """
    将 PDF 转换为可编辑的 Word 文档。

    使用 pdf2docx 重建文字、表格和图片对象；复杂 PDF 的排版可能有轻微差异，
    但不会降级为仅包含页面截图的 DOCX。
    """
    try:
        # 验证文件类型
        upload_name = _safe_upload_name(file.filename, 'document.pdf')
        file_ext = os.path.splitext(upload_name)[1].lower()
        if file_ext != '.pdf':
            raise HTTPException(
                status_code=400,
                detail="仅支持 PDF 文件格式"
            )
        
        # 读取文件内容
        content = await file.read()
        
        temp_dir = tempfile.mkdtemp(prefix="pdf_to_word_")
        temp_pdf_path = os.path.join(temp_dir, upload_name)
        temp_docx_path = os.path.join(temp_dir, f"{Path(upload_name).stem}.docx")
        response_ready = False
        try:
            with open(temp_pdf_path, 'wb') as temp_pdf:
                temp_pdf.write(content)
            from pdf2docx import Converter
            converter = Converter(temp_pdf_path)
            try:
                converter.convert(temp_docx_path)
            finally:
                converter.close()
            if not os.path.exists(temp_docx_path) or os.path.getsize(temp_docx_path) == 0:
                raise RuntimeError('未生成可编辑的 Word 文件')
            response_ready = True
            return FileResponse(
                path=temp_docx_path,
                filename=f"{Path(upload_name).stem}.docx",
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                background=BackgroundTask(_remove_directory, temp_dir),
            )
        finally:
            # FileResponse 接管成功路径的清理；异常路径在此立即释放临时文件。
            if not response_ready:
                _remove_directory(temp_dir)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理文件时出错：{str(e)}")


@app.post("/api/word-to-pdf")
async def convert_word_to_pdf(file: UploadFile = File(...)):
    """将 DOCX 文档转换为 PDF。"""
    upload_name = _safe_upload_name(file.filename, 'document.docx')
    if os.path.splitext(upload_name)[1].lower() != '.docx':
        raise HTTPException(status_code=400, detail="仅支持 DOCX 文件格式")
    temp_dir = tempfile.mkdtemp(prefix="word_to_pdf_")
    source_path = os.path.join(temp_dir, upload_name)
    try:
        with open(source_path, 'wb') as output:
            output.write(await file.read())
        process = _run_soffice(source_path, temp_dir, 'pdf')
        pdf_path = os.path.join(temp_dir, f"{Path(upload_name).stem}.pdf")
        if process.returncode != 0 or not os.path.exists(pdf_path):
            message = process.stderr.strip() or process.stdout.strip() or '未生成 PDF 文件'
            raise HTTPException(status_code=500, detail=f"Word 转 PDF 失败：{message}")
        return FileResponse(
            pdf_path,
            filename=f"{Path(upload_name).stem}.pdf",
            media_type='application/pdf',
            background=BackgroundTask(_remove_directory, temp_dir),
        )
    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=504, detail="Word 转 PDF 超时")
    except HTTPException:
        _remove_directory(temp_dir)
        raise
    except Exception as e:
        _remove_directory(temp_dir)
        raise HTTPException(status_code=500, detail=f"Word 转 PDF 失败：{str(e)}")


@app.post("/api/redact-with-conversion")
async def redact_file_with_conversion(
    file: UploadFile = File(...),
    custom_rules: str = Form("[]")
):
    """
    上传文件并执行初步脱敏（支持 PDF 自动转换为 Word）
    
    处理流程：
    1. 接收用户上传的文件
    2. 如果是 PDF，先转换为 Word 格式（保留格式）
    3. 提取文件文本内容
    4. 使用 Microsoft Presidio 和中国百家姓库检测 PII
    5. 执行初步脱敏（替换为占位符）
    6. 返回脱敏后的文本、映射表和转换后的 Word 文件
    
    参数：
    - file: 上传的文件
    - custom_rules: 本机敏感字段规则 JSON
    
    返回：
    - 脱敏后的文本
    - 映射表（占位符 -> 原始值）
    - 检测到的敏感信息统计
    - 转换后的 Word 文件（如果是 PDF）
    """
    try:
        # 验证文件类型
        allowed_extensions = {'.txt', '.csv', '.json', '.md', '.pdf', '.docx', '.xlsx', '.xls'}
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件格式：{file_ext}。支持的格式：{', '.join(allowed_extensions)}"
            )
        
        # 读取文件内容
        content = await file.read()
        
        # 如果是 PDF，先转换为 Word
        converted_docx_path = None
        if file_ext == '.pdf':
            try:
                # 创建临时文件保存上传的 PDF
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                    temp_pdf.write(content)
                    temp_pdf_path = temp_pdf.name
                
                # 创建临时文件保存转换后的 Word
                converted_docx_path = temp_pdf_path.replace('.pdf', '.docx')
                
                # 使用 pdf2docx 进行转换
                from pdf2docx import Converter
                
                cv = Converter(temp_pdf_path)
                cv.convert(converted_docx_path)
                cv.close()
                
                # 读取转换后的 Word 文件内容
                with open(converted_docx_path, 'rb') as f:
                    docx_content = f.read()
                
                # 使用 Word 文件进行后续处理
                content = docx_content
                file_ext = '.docx'
                
                # 清理临时 PDF 文件
                os.unlink(temp_pdf_path)
                
            except Exception as e:
                # 如果转换失败，继续使用原始 PDF 处理
                print(f"PDF 转换警告：{str(e)}")
        
        # 根据文件类型提取文本
        text_content = desensitization_service.extract_text(content, file_ext)
        document_preview = build_docx_preview(content) if file_ext == '.docx' else None
        
        if not text_content:
            raise HTTPException(
                status_code=400,
                detail="无法从文件中提取文本内容"
            )
        
        # 执行脱敏
        result = desensitization_service.redact_text(text_content, parse_custom_rules(custom_rules))
        
        # 保存脱敏结果（可选）
        task_id = str(uuid.uuid4())
        result["task_id"] = task_id
        result["original_filename"] = file.filename
        result["file_type"] = file_ext
        result["converted_from_pdf"] = file_ext == '.docx' and file.filename.lower().endswith('.pdf')
        if document_preview is not None:
            result["document_preview"] = document_preview
        
        # 保存原始文件
        original_path = os.path.join(UPLOAD_DIR, f"{task_id}_original{file_ext}")
        with open(original_path, "wb") as f:
            f.write(content)
        
        # 保存脱敏结果
        result_path = os.path.join(UPLOAD_DIR, f"{task_id}_result.json")
        with open(result_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        
        # 如果有转换后的 Word 文件，保存它
        if converted_docx_path and os.path.exists(converted_docx_path):
            converted_path = os.path.join(UPLOAD_DIR, f"{task_id}_converted.docx")
            with open(converted_docx_path, 'rb') as src, open(converted_path, 'wb') as dst:
                dst.write(src.read())
            result["converted_docx_path"] = converted_path
            # 清理临时文件
            os.unlink(converted_docx_path)
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理文件时出错：{str(e)}")


@app.post("/api/detect")
async def detect_sensitive_content(file: UploadFile = File(...)):
    """
    检测文件中的敏感信息（不执行脱敏）
    
    支持的文件格式：
    - 文本文件：TXT, CSV, JSON, MD
    - PDF 文件
    - Word 文档：DOCX
    - Excel 文件：XLSX, XLS
    
    返回：
    - 检测到的敏感信息列表
    - 敏感信息类型（姓名、手机号、身份证、银行卡等）
    - 位置信息（行号、列号、偏移量）
    """
    try:
        # 验证文件类型
        allowed_extensions = {'.txt', '.csv', '.json', '.md', '.pdf', '.docx', '.xlsx', '.xls'}
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件格式：{file_ext}。支持的格式：{', '.join(allowed_extensions)}"
            )
        
        # 读取文件内容
        content = await file.read()
        
        # 根据文件类型提取文本
        text_content = desensitization_service.extract_text(content, file_ext)
        document_preview = build_docx_preview(content) if file_ext == '.docx' else None
        
        if not text_content:
            raise HTTPException(
                status_code=400,
                detail="无法从文件中提取文本内容"
            )
        
        # 检测敏感信息
        detections = desensitization_service.detect_sensitive_info(text_content)
        
        return {
            "success": True,
            "filename": file.filename,
            "file_type": file_ext,
            "text_length": len(text_content),
            "detections": detections,
            "detection_count": len(detections)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理文件时出错：{str(e)}")


@app.post("/api/redact")
async def redact_file(
    file: UploadFile = File(...),
    custom_rules: str = Form("[]")
):
    """
    上传文件并执行初步脱敏
    
    处理流程：
    1. 接收用户上传的文件
    2. 提取文件文本内容
    3. 使用 Microsoft Presidio 和中国百家姓库检测 PII
    4. 执行初步脱敏（替换为占位符）
    5. 返回脱敏后的文本和映射表
    
    参数：
    - file: 上传的文件
    - custom_rules: 本机敏感字段规则 JSON
    
    返回：
    - 脱敏后的文本
    - 映射表（占位符 -> 原始值）
    - 检测到的敏感信息统计
    """
    try:
        # 验证文件类型
        allowed_extensions = {'.txt', '.csv', '.json', '.md', '.pdf', '.docx', '.xlsx', '.xls'}
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件格式：{file_ext}。支持的格式：{', '.join(allowed_extensions)}"
            )
        
        # 读取文件内容
        content = await file.read()
        
        # 根据文件类型提取文本
        text_content = desensitization_service.extract_text(content, file_ext)
        document_preview = build_docx_preview(content) if file_ext == '.docx' else None
        
        if not text_content:
            raise HTTPException(
                status_code=400,
                detail="无法从文件中提取文本内容"
            )
        
        # 执行脱敏
        result = desensitization_service.redact_text(text_content, parse_custom_rules(custom_rules))
        
        # 保存脱敏结果（可选）
        task_id = str(uuid.uuid4())
        result["task_id"] = task_id
        result["original_filename"] = file.filename
        result["file_type"] = file_ext
        if document_preview is not None:
            result["document_preview"] = document_preview
        
        # 保存原始文件
        original_path = os.path.join(UPLOAD_DIR, f"{task_id}_original{file_ext}")
        with open(original_path, "wb") as f:
            f.write(content)
        
        # 保存脱敏结果
        result_path = os.path.join(UPLOAD_DIR, f"{task_id}_result.json")
        with open(result_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理文件时出错：{str(e)}")


@app.post("/api/redact-preserving-format")
async def redact_preserving_format(
    file: UploadFile = File(...),
    mappings: str = Form(...),
):
    """根据人工复核后的映射表输出保留原始结构的 DOCX/XLSX；PDF 输出 DOCX。"""
    try:
        file_ext = os.path.splitext(file.filename or '')[1].lower()
        mapping_data = json.loads(mappings)
        mapping_list = mapping_data.get('mappings', []) if isinstance(mapping_data, dict) else mapping_data
        if file_ext not in {'.docx', '.xlsx', '.pdf'}:
            raise HTTPException(status_code=400, detail='仅支持 DOCX、XLSX 和 PDF 的保留格式脱敏输出')
        output, output_ext = redact_file_preserving_format(await file.read(), file_ext, mapping_list)
        with tempfile.NamedTemporaryFile(suffix=output_ext, delete=False) as temp_file:
            temp_file.write(output)
            temp_path = temp_file.name
        filename = f"redacted_{Path(file.filename or 'document').stem}{output_ext}"
        media_type = (
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            if output_ext == '.docx' else 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        return FileResponse(
            path=temp_path, filename=filename, media_type=media_type,
            background=BackgroundTask(os.unlink, temp_path),
        )
    except HTTPException:
        raise
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail='映射表格式错误')
    except Exception as e:
        raise HTTPException(status_code=500, detail=f'保留格式脱敏失败：{str(e)}')


@app.post("/api/restore")
async def restore_file(
    redacted_file: UploadFile = File(...),
    mapping_file: UploadFile = File(...)
):
    """
    还原脱敏文件
    
    处理流程：
    1. 接收脱敏文件和映射表
    2. 根据映射表还原原始内容
    4. 返回还原后的文件
    
    参数：
    - redacted_file: 脱敏后的文件
    - mapping_file: 映射表 JSON 文件
    
    返回：
    - 还原后的文本
    - 还原统计信息
    """
    try:
        # 读取映射表
        mapping_content = await mapping_file.read()
        mapping_data = json.loads(mapping_content)
        
        # 读取脱敏文件
        redacted_content = await redacted_file.read()
        file_ext = os.path.splitext(redacted_file.filename)[1].lower()
        
        # 提取文本
        redacted_text = desensitization_service.extract_text(redacted_content, file_ext)
        
        if not redacted_text:
            raise HTTPException(
                status_code=400,
                detail="无法从脱敏文件中提取文本内容"
            )
        
        # 执行还原
        result = desensitization_service.restore_text(redacted_text, mapping_data)
        
        return result
        
    except HTTPException:
        raise
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="映射表文件格式错误，请上传有效的 JSON 文件")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"还原文件时出错：{str(e)}")


@app.post("/api/text-to-word")
async def convert_text_to_word(
    text: str = Form(...),
    filename: str = Form("document.docx"),
    redaction_notice: bool = Form(False),
):
    """
    将文本转换为 Word 文档
    """
    try:
        from docx import Document
        
        doc = Document()
        if redaction_notice:
            _add_docx_redaction_notice(doc)
        paragraphs = text.split('\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text)
            else:
                doc.add_paragraph('')
        
        temp_path = f"/tmp/{filename}"
        doc.save(temp_path)
        
        return FileResponse(
            path=temp_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"转换失败：{str(e)}")


@app.post("/api/text-to-excel")
async def convert_text_to_excel(
    text: str = Form(...),
    filename: str = Form("document.xlsx"),
    redaction_notice: bool = Form(False),
):
    """
    将文本转换为 Excel 文件（每行一个单元格）
    """
    try:
        from openpyxl import Workbook
        
        # 按行分割文本
        lines = text.split('\n')
        
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = '内容'
        worksheet.append(['内容'])
        for line in lines:
            worksheet.append([line])
        if redaction_notice:
            _add_xlsx_redaction_notice(workbook)
        temp_path = f"/tmp/{filename}"
        workbook.save(temp_path)
        
        return FileResponse(
            path=temp_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"转换失败：{str(e)}")


@app.post("/api/text-to-markdown")
async def convert_text_to_markdown(
    text: str = Form(...),
    filename: str = Form("document.md")
):
    """
    将文本转换为 Markdown 文件
    """
    try:
        # 直接将文本保存为 Markdown
        temp_path = f"/tmp/{filename}"
        with open(temp_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        return FileResponse(
            path=temp_path,
            filename=filename,
            media_type='text/markdown'
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"转换失败：{str(e)}")


@app.post("/api/text-to-txt")
async def convert_text_to_txt(
    text: str = Form(...),
    filename: str = Form("document.txt")
):
    """
    将文本转换为 TXT 文件
    """
    try:
        temp_path = f"/tmp/{filename}"
        with open(temp_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        return FileResponse(
            path=temp_path,
            filename=filename,
            media_type='text/plain'
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"转换失败：{str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
